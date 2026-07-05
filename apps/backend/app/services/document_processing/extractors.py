import abc
import csv
import json
import logging
from io import BytesIO, StringIO
import re
import html
from typing import Dict, Type

# Attempt imports of optional document processing libraries
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx
except ImportError:
    docx = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

logger = logging.getLogger("app.services.document_processing.extractors")


class BaseExtractor(abc.ABC):
    """
    Abstract Base Class for all document content extractors.
    """
    @abc.abstractmethod
    def extract(self, content: bytes) -> str:
        """
        Extracts and returns plain text from raw file bytes.
        """
        pass

    def _clean(self, text: str) -> str:
        """
        Helper to clean up whitespaces, control characters, and redundant blank lines.
        """
        text = text.replace("\x00", "")
        text = re.sub(r"\r\n", "\n", text)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()


class PlainTextExtractor(BaseExtractor):
    def extract(self, content: bytes) -> str:
        text = content.decode("utf-8", errors="replace")
        return self._clean(text)


class HTMLExtractor(BaseExtractor):
    def extract(self, content: bytes) -> str:
        text = content.decode("utf-8", errors="replace")
        # Strip script and style tags completely
        text = re.sub(r"<script[^>]*>.*?</script>", "", text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
        # Strip all HTML tags
        text = re.sub(r"<[^>]+>", " ", text)
        text = html.unescape(text)
        return self._clean(text)


class CSVExtractor(BaseExtractor):
    def extract(self, content: bytes) -> str:
        text = content.decode("utf-8", errors="replace")
        reader = csv.reader(StringIO(text))
        rows = []
        for row in reader:
            if row:
                rows.append(" | ".join(row))
        return self._clean("\n".join(rows))


class JSONExtractor(BaseExtractor):
    def extract(self, content: bytes) -> str:
        text = content.decode("utf-8", errors="replace")
        try:
            obj = json.loads(text)
            return self._clean(json.dumps(obj, indent=2, ensure_ascii=False))
        except json.JSONDecodeError:
            return self._clean(text)


class PDFExtractor(BaseExtractor):
    def extract(self, content: bytes) -> str:
        if not PdfReader:
            logger.warning("pypdf is not installed. Returning placeholder.")
            return "[PDF content - pypdf not available]"
        
        try:
            reader = PdfReader(BytesIO(content))
            pages_text = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages_text.append(f"--- Page {i+1} ---\n{text}")
            return self._clean("\n\n".join(pages_text))
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            raise ValueError(f"Failed to parse PDF content: {e}")


class DOCXExtractor(BaseExtractor):
    def extract(self, content: bytes) -> str:
        if not docx:
            logger.warning("python-docx is not installed. Returning placeholder.")
            return "[DOCX content - python-docx not available]"

        try:
            doc = docx.Document(BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))
            
            return self._clean("\n\n".join(paragraphs))
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            raise ValueError(f"Failed to parse DOCX content: {e}")


class ExcelExtractor(BaseExtractor):
    def extract(self, content: bytes) -> str:
        if not openpyxl:
            logger.warning("openpyxl is not installed. Returning placeholder.")
            return "[Excel content - openpyxl not available]"

        try:
            wb = openpyxl.load_workbook(BytesIO(content), data_only=True, read_only=True)
            sheet_texts = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = []
                for row in ws.iter_rows(values_only=True):
                    # Filter out empty rows, convert cells to string
                    cells = [str(cell).strip() if cell is not None else "" for cell in row]
                    if any(cells):
                        rows.append(" | ".join(cells))
                if rows:
                    sheet_texts.append(f"--- Sheet: {sheet_name} ---\n" + "\n".join(rows))
            return self._clean("\n\n".join(sheet_texts))
        except Exception as e:
            logger.error(f"Error extracting Excel: {e}")
            raise ValueError(f"Failed to parse Excel content: {e}")


class ImagePlaceholderExtractor(BaseExtractor):
    def extract(self, content: bytes) -> str:
        # Standard placeholder for future OCR module inclusion
        return "[Image file uploaded successfully. Metadata cached. OCR processing pending.]"


class DocumentExtractorFactory:
    """
    Factory to resolve the correct text extractor strategy based on MIME type or file extension.
    """
    _mapping: Dict[str, Type[BaseExtractor]] = {
        "text/plain": PlainTextExtractor,
        "text/markdown": PlainTextExtractor,
        "text/html": HTMLExtractor,
        "text/csv": CSVExtractor,
        "application/json": JSONExtractor,
        "application/pdf": PDFExtractor,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DOCXExtractor,
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ExcelExtractor,
        "application/vnd.ms-excel": ExcelExtractor,
        "image/png": ImagePlaceholderExtractor,
        "image/jpeg": ImagePlaceholderExtractor,
        "image/gif": ImagePlaceholderExtractor,
        "image/webp": ImagePlaceholderExtractor,
    }

    _extension_mapping: Dict[str, Type[BaseExtractor]] = {
        "txt": PlainTextExtractor,
        "md": PlainTextExtractor,
        "rst": PlainTextExtractor,
        "html": HTMLExtractor,
        "csv": CSVExtractor,
        "json": JSONExtractor,
        "pdf": PDFExtractor,
        "docx": DOCXExtractor,
        "xlsx": ExcelExtractor,
        "xls": ExcelExtractor,
        "png": ImagePlaceholderExtractor,
        "jpg": ImagePlaceholderExtractor,
        "jpeg": ImagePlaceholderExtractor,
        "webp": ImagePlaceholderExtractor,
    }

    @classmethod
    def get_extractor(cls, mime_type: str, filename: str) -> BaseExtractor:
        # Resolve by extension first to override broad MIME types
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if ext in cls._extension_mapping:
            return cls._extension_mapping[ext]()

        # Resolve by MIME type
        if mime_type in cls._mapping:
            return cls._mapping[mime_type]()

        # Fallback to plain text decode or placeholder
        return PlainTextExtractor()
